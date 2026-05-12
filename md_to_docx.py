"""Convert DOCUMENTATION.md → DOCUMENTATION.docx (conservative Word-compatible output)."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


SRC = Path(__file__).parent / "DOCUMENTATION.md"
DST = Path(__file__).parent / "DOCUMENTATION.docx"


_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_BOLD = re.compile(r"\*\*([^*]+)\*\*")
_CODE = re.compile(r"`([^`]+)`")


def add_inline(paragraph, text: str) -> None:
    """Add a markdown-flavored line to a paragraph, handling bold/code/links."""
    i = 0
    while i < len(text):
        m_bold = _BOLD.search(text, i)
        m_code = _CODE.search(text, i)
        m_link = _LINK.search(text, i)
        candidates = [m for m in (m_bold, m_code, m_link) if m]
        if not candidates:
            paragraph.add_run(text[i:])
            return
        first = min(candidates, key=lambda m: m.start())
        if first.start() > i:
            paragraph.add_run(text[i:first.start()])
        if first is m_bold:
            run = paragraph.add_run(first.group(1))
            run.bold = True
        elif first is m_code:
            run = paragraph.add_run(first.group(1))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            run = paragraph.add_run(first.group(1))
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            run.underline = True
        i = first.end()


def add_code_block(doc, code_lines: list[str]) -> None:
    """Add a code block. Each line is its own paragraph for safe Word rendering."""
    for line in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def convert() -> None:
    text = SRC.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Glimmora Relocate — Technical Documentation")
    run.bold = True
    run.font.size = Pt(22)
    doc.add_paragraph()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # Skip the H1 we render as the title
        if stripped.startswith("# Glimmora Relocate"):
            i += 1
            continue

        # Fenced code block
        if stripped.startswith("```"):
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            add_code_block(doc, code_lines)
            continue

        # Horizontal rule
        if stripped in ("---", "***"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("─" * 40)
            i += 1
            continue

        # Table (header + separator + body rows)
        if (
            stripped.startswith("|")
            and i + 1 < len(lines)
            and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip())
        ):
            header_cells = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
            table.style = "Table Grid"
            for j, h in enumerate(header_cells):
                cell = table.rows[0].cells[j]
                cell.text = ""
                add_inline(cell.paragraphs[0], h)
                for r in cell.paragraphs[0].runs:
                    r.bold = True
            for r_idx, row in enumerate(rows, start=1):
                for j, cell_val in enumerate(row):
                    if j >= len(table.rows[r_idx].cells):
                        break
                    cell = table.rows[r_idx].cells[j]
                    cell.text = ""
                    add_inline(cell.paragraphs[0], cell_val)
            doc.add_paragraph()
            continue

        # Headings
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        # Bullet list
        if re.match(r"^[\-\*] ", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s", "", stripped))
            i += 1
            continue

        # Blockquote
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(stripped[2:])
            run.italic = True
            i += 1
            continue

        # Blank line
        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    doc.save(DST)
    print(f"wrote {DST} ({DST.stat().st_size} bytes)")


if __name__ == "__main__":
    convert()
